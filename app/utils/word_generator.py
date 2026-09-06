from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import re
from datetime import datetime
from typing import Dict, Any, List


class WordGenerator:
    """Word文档生成器（支持 Markdown 语法渲染）"""

    def __init__(self):
        self.doc = Document()
        self._setup_chinese_fonts()

    def _setup_chinese_fonts(self):
        """配置中文字体支持"""
        styles = self.doc.styles
        for style in styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                try:
                    style.font.name = 'SimSun'
                    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                except Exception:
                    pass

    # ---------------------------------------------------------------
    # 字体辅助
    # ---------------------------------------------------------------
    def _apply_font(self, run, size: int = 11, bold: bool = False,
                    italic: bool = False, name: str = 'SimSun', color=None):
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.name = name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
        if color is not None:
            run.font.color.rgb = color

    def _add_inline_md(self, para, text: str):
        """把行内 Markdown（**加粗** *斜体* `代码`）渲染为带格式 run"""
        if text is None:
            return
        pattern = re.compile(r'(\*\*.+?\*\*|\*.+?\*|`.+?`)')
        pos = 0
        for m in pattern.finditer(text):
            if m.start() > pos:
                r = para.add_run(text[pos:m.start()])
                self._apply_font(r)
            tok = m.group(1)
            if tok.startswith('**'):
                r = para.add_run(tok[2:-2])
                self._apply_font(r, bold=True)
            elif tok.startswith('`'):
                r = para.add_run(tok[1:-1])
                self._apply_font(r, name='Consolas')
            else:
                r = para.add_run(tok[1:-1])
                self._apply_font(r, italic=True)
            pos = m.end()
        if pos < len(text):
            r = para.add_run(text[pos:])
            self._apply_font(r)

    # ---------------------------------------------------------------
    # 封面 / 目录
    # ---------------------------------------------------------------
    def generate_cover(self, title: str, subtitle: str = None,
                       date: str = None, customer: str = None):
        """生成封面页"""
        self.doc.add_paragraph()
        self.doc.add_paragraph()

        title_para = self.doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.name = 'SimHei'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_paragraph()

        if subtitle:
            sub_para = self.doc.add_paragraph()
            sub_run = sub_para.add_run(subtitle)
            sub_run.font.size = Pt(18)
            sub_run.font.name = 'SimSun'
            sub_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_paragraph()
        self.doc.add_paragraph()

        if customer:
            cust_para = self.doc.add_paragraph()
            cust_run = cust_para.add_run(f"客户：{customer}")
            cust_run.font.size = Pt(13)
            cust_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if date:
            date_para = self.doc.add_paragraph()
            date_run = date_para.add_run(f"生成日期：{date}")
            date_run.font.size = Pt(12)
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 品牌标识
        self.doc.add_paragraph()
        brand_para = self.doc.add_paragraph()
        brand_run = brand_para.add_run("华为云 · 解决方案智能匹配自动生成")
        brand_run.font.size = Pt(11)
        brand_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        brand_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_page_break()

    def _strip_number(self, title: str) -> str:
        """去掉标题前缀的 '1. ' / '1、' 之类的编号（由目录/正文统一重新编号）"""
        return re.sub(r'^\s*\d+[.、]\s*', '', title or '').strip()

    def generate_toc(self, chapters: List[Dict]):
        """生成目录页（序号由本方法统一生成，避免与标题内编号重复）"""
        toc_title = self.doc.add_paragraph()
        toc_run = toc_title.add_run("目 录")
        toc_run.font.size = Pt(18)
        toc_run.font.bold = True
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_paragraph()

        for idx, chapter in enumerate(chapters, 1):
            toc_item = self.doc.add_paragraph()
            clean = self._strip_number(chapter.get('title', '未命名章节'))
            toc_item.add_run(f"{idx}. {clean}")

        self.doc.add_page_break()

    # ---------------------------------------------------------------
    # Markdown 块渲染
    # ---------------------------------------------------------------
    def _add_horizontal_rule(self):
        p = self.doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'BBBBBB')
        pbdr.append(bottom)
        pPr.append(pbdr)

    def _split_row(self, line: str) -> List[str]:
        s = line.strip()
        if s.startswith('|'):
            s = s[1:]
        if s.endswith('|'):
            s = s[:-1]
        return [c.strip() for c in s.split('|')]

    def _set_cell(self, cell, text: str, bold: bool = False):
        cell.text = ""
        para = cell.paragraphs[0]
        self._add_inline_md(para, text)
        # 整单元格加粗（表头）
        for run in para.runs:
            run.font.bold = bold

    def _add_table(self, rows: List[str]):
        if len(rows) < 1:
            return
        # 2026-09-06 修：调用方 _render_markdown 收集 block 时从分隔行之后开始
        # （不含分隔行），因此数据从 rows[1] 起。旧逻辑误以为 rows[1] 是分隔行、
        # 从 rows[2] 取数据 → 每张表格的第一行数据被静默丢弃（成本附表首行 SKU 消失）。
        # 防御：过滤任何形如 | :--- | 的分隔行。
        header = self._split_row(rows[0])
        is_sep = lambda r: bool(re.match(r'^\|[\s\-:|]+\|$', r.strip()))
        data = [self._split_row(r) for r in rows[1:] if not is_sep(r)]
        ncol = max(len(header), 1)
        table = self.doc.add_table(rows=1, cols=ncol)
        try:
            table.style = 'Table Grid'
        except Exception:
            pass
        hdr_cells = table.rows[0].cells
        for c in range(ncol):
            self._set_cell(hdr_cells[c], header[c] if c < len(header) else '', bold=True)
        for row in data:
            cells = table.add_row().cells
            for c in range(ncol):
                self._set_cell(cells[c], row[c] if c < len(row) else '')
        self.doc.add_paragraph()

    def _add_list_item(self, text: str, numbered: bool = False):
        # 不用 'List Bullet'/'List Number' 内置样式（带圆点/编号符号，与网页端无圆点风格不一致）
        # 改为普通段落 + 左缩进，保持列表层级感但无任何前缀符号
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Pt(18)
        self._add_inline_md(para, text)

    def _render_markdown(self, content: str):
        """把 Markdown 文本块渲染为 Word 元素（段落/列表/表格/分隔符）"""
        lines = (content or "").split("\n")
        i = 0
        n = len(lines)
        while i < n:
            line = lines[i]
            stripped = line.strip()

            # 水平分隔符
            if re.match(r'^(\*\*\*|___|---)$', stripped):
                self._add_horizontal_rule()
                i += 1
                continue

            # 表格块：当前行以 | 开头，向后查找分隔行（跳过空行，最多看3行）
            if stripped.startswith("|"):
                sep_idx = -1
                for look in range(i + 1, min(i + 4, n)):
                    ls = lines[look].strip()
                    if re.match(r'^\|[\s\-:|]{3,}\|$', ls):
                        sep_idx = look
                        break
                    if ls != "" and not ls.startswith("|"):
                        break
                if sep_idx != -1:
                    block = [stripped]
                    j = sep_idx + 1
                    while j < n and lines[j].strip().startswith("|"):
                        block.append(lines[j].strip())
                        j += 1
                    self._add_table(block)
                    i = j
                    continue

            # 无序列表
            m = re.match(r'^([\*\-])\s+(.*)', stripped)
            if m:
                self._add_list_item(m.group(2), numbered=False)
                i += 1
                continue

            # 有序列表
            m = re.match(r'^(\d+)\.\s+(.*)', stripped)
            if m:
                self._add_list_item(m.group(2), numbered=True)
                i += 1
                continue

            # 空行
            if stripped == "":
                i += 1
                continue

            # 普通段落
            para = self.doc.add_paragraph()
            self._add_inline_md(para, stripped)
            i += 1

    # ---------------------------------------------------------------
    # 章节
    # ---------------------------------------------------------------
    def generate_chapter(self, title: str, content: str, level: int = 1):
        """生成章节内容（content 为 Markdown，内部转换为 Word 格式）"""
        self.doc.add_heading(title, level=level)
        self._render_markdown(content)

    def generate_report(self, report_data: Dict[str, Any]) -> Document:
        """生成完整报告"""
        title = report_data.get('title', '华为云解决方案报告')
        subtitle = report_data.get('subtitle', '')
        date = report_data.get('create_date', datetime.now().strftime('%Y-%m-%d'))
        customer = report_data.get('customer_name', '')
        chapters = report_data.get('chapters', [])
        appendix = report_data.get('appendix', [])

        self.generate_cover(title, subtitle, date, customer)

        if chapters:
            self.generate_toc(chapters)

        for idx, chapter in enumerate(chapters, 1):
            ch_title = chapter.get('title', '')
            ch_content = chapter.get('content', '')
            display_title = f"{idx}. {self._strip_number(ch_title)}"
            self.generate_chapter(display_title, ch_content)

            sub_sections = chapter.get('sections', [])
            for sub in sub_sections:
                sub_title = sub.get('title', '')
                sub_content = sub.get('content', '')
                self.generate_chapter(sub_title, sub_content, level=2)

        if appendix:
            self.doc.add_page_break()
            self.doc.add_heading("附 录", level=1)
            for item in appendix:
                app_title = item.get('title', '')
                app_content = item.get('content', '')
                self.generate_chapter(app_title, app_content, level=2)

        return self.doc

    def save(self, file_path: str):
        """保存文档"""
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        self.doc.save(file_path)
